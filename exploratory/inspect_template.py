
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
path = ROOT / "Lab Notes Template.docx"
try:
    doc = Document(path)
    print("Document Structure:")
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"Paragraph: {p.text}")
            print(f"  Style: {p.style.name}")
    
    print("\nTables:")
    for t in doc.tables:
        print("  Table found")
        for r in t.rows:
            row_text = [c.text for c in r.cells]
            print(f"  Row: {row_text}")

except Exception as e:
    print(f"Error reading template: {e}")
